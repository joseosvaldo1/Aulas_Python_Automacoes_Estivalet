"""
Documentos Word (Parte III):


* Criando Arquivos Docx:
- Heading:

"""
import docx
import sys

# Definir o encoding padrão para utf-8
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document()

doc.add_heading('Header 0 - Título', 0)
doc.add_heading('Header 1 - Subtítulo', 1)
doc.add_heading('Header 2 - Texto 1', 2)
doc.add_heading('Header 3 - Texto 2', 3)
doc.add_heading('Header 4 - Notas de Rodapé', 4)
path_4 = r'.\word\headings.docx'
doc.save(path_4)
