"""
Documentos Word (Parte II):


* Mudando Estilo de Texto:

"""
import docx
import sys

# Definir o encoding padrão para utf-8
sys.stdout.reconfigure(encoding='utf-8')
path = r'.\word\demo.docx'
doc = docx.Document(path)
print(f"doc.paragraphs[0].style = {doc.paragraphs[0].style}")
print(" ")
print("Alterando o estilo de 'doc.paragraph[0].style':")
doc.paragraphs[0].style = 'Normal'
print(f"doc.paragraphs[0].style = {doc.paragraphs[0].style}")
path_2 = r'.\word\demo2.docx'
doc.save(path_2)

