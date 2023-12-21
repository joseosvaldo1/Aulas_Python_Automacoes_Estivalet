"""
Documentos Word (Parte I):

* Lendo Documentos .docx
- O Python pode criar e modificar documentos Word, que têm 
extensão de arquivo .docx, com o módulo python-docx. 
- A documentação completa do Python-Docx está disponível
em https://python-docx.readthedocs.org/.
- Comparados aos arquivos em formato texto simples, os arquivos .docx
contêm muitas estruturas. Essas estruturas são representadas por 'três 
tipos de dados' diferentes no Python-Docx. Em âmbito geral, um objeto 
'Document representa o documento completo'. O objeto Document contém 
uma lista de objetos 'Paragraph para os parágrafos do documento'. 
(Um novo parágrafo começa sempre que o usuário tecla enter ou return 
enquanto está digitando um documento Word.) 'Cada um desses objetos 
Paragraph contém uma lista de um ou mais objetos Run'.

Exemplo: A plain paragraph with some        bold           and some             italic.
|                  run               | run             |     run          |       run        |
|             style_1 (normal)       | style_2 (bold)  | style_3 (normal) | style_4 (normal) |

- O texto em um documento Word é mais que apenas uma string. 
Esse texto tem fonte, tamanho, cor e outras informações de 
estilização associadas a ele. Um estilo (style) em Word é uma 
coleção desses atributos. Um objeto Run é uma porção contígua 
de texto que apresenta o mesmo estilo. Um novo objeto Run é 
necessário sempre que o estilo do texto mudar.


"""
import docx
import sys

# Definir o encoding padrão para utf-8
sys.stdout.reconfigure(encoding='utf-8')
path = r'.\word\demo.docx'
doc = docx.Document(path)
# Determinando a quantidade de parágrafos por meio do método 'len()':
print(f"Número de parágrafos: {len(doc.paragraphs)}")
print(" ")
# Extraindo textos de um dado parágrafo:
print(f"Parágrafo 1: {doc.paragraphs[0].text}")
print(f"Parágrafo 2: {doc.paragraphs[1].text}")
print(" ")
for i in range(len(doc.paragraphs)):
    print(f"Parágrafo {i+1}: {doc.paragraphs[i].text}")

print(" ")
# Determinando a quantidade de objetos 'run' do parágrafo 2:
print(f"Runs do Parágrafo 2: {len(doc.paragraphs[1].runs)}")
print(" ")
# Determinando o primeiro ran do segundo parágrafo:
print("Runs do segundo parágrafo:")

for i in range(len(doc.paragraphs[1].runs)):
    print(f"doc.paragraphs[1].runs[{i}]: {doc.paragraphs[1].runs[i].text}")
print(" ")
# Pegando todo o texto independente dos run (estilos):
full_text = []
for parag in doc.paragraphs:
    full_text.append(parag.text)

print("Texto Completo:")
print('\n'.join(full_text))