"""
Documentos Word (Parte III):


* Criando Arquivos Docx:

"""
import docx
import sys

# Definir o encoding padrão para utf-8
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document()
# Adicionando um parágrafo com o método 'add_paragraph()' 
# ao documento todo: 
doc.add_paragraph('Hello World!')
# Adicionando um parágrafo com o método 'add_paragraph()' 
# a um objeto paragraph:
parag_obj_1 = doc.add_paragraph('This is the second paragraph!')
parag_obj_2 = doc.add_paragraph('This is the third paragraph!')

# Adicionando outro texto com o método 'add_run()' a um 
# objeto paragraph:
parag_obj_1.add_run(' This text is being added to the second paragraph.')

path_3 = r'.\word\mutiple_paragraphs.docx'
doc.save(path_3)

