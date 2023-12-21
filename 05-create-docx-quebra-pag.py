"""
Documentos Word (Parte III):


* Criando Arquivos Docx:

- Quebra de Página:

"""
import docx
import sys

# Definir o encoding padrão para utf-8
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document()
doc.add_paragraph("This is the first page.")
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph("This is the second page.")
path_5 = r'.\word\picture.docx'
path_6 = r'.\word\python.png'
doc.add_picture(path_6, width=docx.shared.Cm(8), height=docx.shared.Cm(4))
doc.save(path_5)